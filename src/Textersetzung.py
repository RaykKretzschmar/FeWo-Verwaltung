from docx import Document
import datetime

def replace_text(replacements, save_path, doc_path="Rechnungsvorlage.docx"):
    doc = Document(doc_path)

    # Go through each paragraph in the document
    for p in doc.paragraphs:
        full_text = p.text
        for old_text, new_text in replacements.items():
            full_text = full_text.replace(old_text, new_text)
        # Clear the paragraph and replace with new text
        p.clear()
        p.add_run(full_text)
    
    # Go through each table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text = cell.text
                for old_text, new_text in replacements.items():
                    full_text = full_text.replace(old_text, new_text)
                # Clear the cell and replace with new text
                cell.text = full_text

    if not save_path:
        datum = datetime.datetime.now()
        tag = datum.strftime("%d")
        monat = datum.strftime("%m")
        jahr = datum.strftime("%Y")
        save_path = f"Rechnung{tag}{monat}{jahr}.docx"
    
    # Save the document
    doc.save(save_path.replace('/', ''))
    
    