from django.shortcuts import render, redirect, get_object_or_404
from .models import Invoice
from .forms import InvoiceForm
from customers.models import Customer
from properties.models import Property
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.conf import settings
from django.core.files.base import File
import os
from docx import Document
# import pythoncom

@login_required
def invoice_list(request):
    invoices = Invoice.objects.filter(user=request.user).order_by("-date")
    return render(request, "invoices/invoice_list.html", {"invoices": invoices})

@login_required
def invoice_create(request):
    if request.method == "POST":
        form = InvoiceForm(request.POST)
        # Filter properties in form if needed, but for now just handle save
        if form.is_valid():
            invoice = form.save(commit=False)
            invoice.user = request.user
            invoice.save()
            try:
                generate_invoice_documents(invoice)
                messages.success(request, "Rechnung erfolgreich erstellt.")
            except Exception as e:
                messages.error(request, f"Fehler bei der Dokumentenerstellung: {e}")
            return redirect("invoice_list")
    else:
        form = InvoiceForm()
        # Ideally filter foreign keys in form here, but standard form usage might show all. 
        # For strict isolation, we should update forms too, but let's start with view logic.
        form.fields['rental_property'].queryset = Property.objects.filter(user=request.user)
        form.fields['customer'].queryset = Customer.objects.filter(user=request.user)

    return render(request, "invoices/invoice_form.html", {"form": form, "title": "Neue Rechnung"})

@login_required
def invoice_create_for_customer(request, customer_id=None):
    customer = None
    if customer_id:
        customer = get_object_or_404(Customer, id=customer_id, user=request.user)

    if request.method == "POST":
        form = InvoiceForm(request.POST)
        if form.is_valid():
            invoice = form.save(commit=False)
            invoice.user = request.user
            if customer:
                invoice.customer = customer
            invoice.save()
            try:
                generate_invoice_documents(invoice)
                messages.success(request, "Rechnung erfolgreich erstellt.")
            except Exception as e:
                messages.error(request, f"Fehler bei der Dokumentenerstellung: {e}")
            return redirect("invoice_list")
    else:
        initial = {}
        if customer:
            initial["customer"] = customer
        form = InvoiceForm(initial=initial)
        form.fields['rental_property'].queryset = Property.objects.filter(user=request.user)
        form.fields['customer'].queryset = Customer.objects.filter(user=request.user)

    return render(
        request, "invoices/invoice_form.html", {"form": form, "customer": customer, "title": "Rechnung erstellen"}
    )

@login_required
def invoice_update(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk, user=request.user)
    if request.method == "POST":
        form = InvoiceForm(request.POST, instance=invoice)
        if form.is_valid():
            invoice = form.save(commit=False)
            invoice.user = request.user
            invoice.save()
            try:
                generate_invoice_documents(invoice)
                messages.success(request, "Rechnung erfolgreich aktualisiert.")
            except Exception as e:
                messages.error(request, f"Fehler bei der Dokumentenerstellung: {e}")
            return redirect("invoice_list") # or similar
    else:
        form = InvoiceForm(instance=invoice)
        # Filter foreign keys just like in create
        form.fields['rental_property'].queryset = Property.objects.filter(user=request.user)
        form.fields['customer'].queryset = Customer.objects.filter(user=request.user)

    return render(
        request, "invoices/invoice_form.html", {"form": form, "title": f"Rechnung {invoice.invoice_number} bearbeiten"}
    )

@login_required
def invoice_delete(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk, user=request.user)
    
    if request.method == "POST":
        # Delete associated files if they exist
        if invoice.pdf_file:
            if os.path.exists(invoice.pdf_file.path):
                os.remove(invoice.pdf_file.path)
        if invoice.docx_file:
            if os.path.exists(invoice.docx_file.path):
                os.remove(invoice.docx_file.path)
        
        invoice_number = invoice.invoice_number
        invoice.delete()
        messages.success(request, f"Rechnung {invoice_number} wurde erfolgreich gelöscht.")
        return redirect("invoice_list")
    
    return render(request, "invoices/invoice_confirm_delete.html", {"invoice": invoice})

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        for run in p.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(str(old_text), str(new_text))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(str(old_text), str(new_text))

def generate_invoice_documents(invoice: Invoice):
    # Path to template
    # The template is in the root directory (two levels up from BASE_DIR which is fewo_web/fewo)
    template_path = settings.BASE_DIR.parent.parent / "Rechnungsvorlage.docx"
    if not os.path.exists(template_path):
        # Fallback if not found in root, try static or media
        raise FileNotFoundError(f"Template not found at {template_path}")

    doc = Document(template_path)

    # Prepare data
    replacements = {
        "Rechnungsnummer": invoice.invoice_number,
        "Datum": invoice.date.strftime("%d.%m.%Y"),
        "Anreisedatum": invoice.arrival_date.strftime("%d.%m.%Y"),
        "Abreisedatum": invoice.departure_date.strftime("%d.%m.%Y"),
        "NdFeWo": invoice.rental_property.name if invoice.rental_property else "",
        "PpN": f"{invoice.price_per_night:.2f}".replace(".", ","),
        "AnzahlDerÜbernachtungen": str(invoice.nights),
        "ÜNKosten": f"{invoice.lodging_total:.2f}".replace(".", ","),
        "GesamtBetrag": f"{invoice.total_price:.2f}".replace(".", ","),
        "MwstBetrag": f"{invoice.tax_amount:.2f}".replace(".", ","),
        "NettoBetrag": f"{invoice.net_amount:.2f}".replace(".", ","),
        # Customer data
        "Vorname": invoice.customer.first_name,
        "Nachname": invoice.customer.last_name,
        "Stadt": invoice.customer.city,
        "PLZ": invoice.customer.postal_code,
        "Straße": invoice.customer.street,
        "Hausnummer": invoice.customer.house_number,
        "Kundennummer": invoice.customer.customer_number,
    }
    
    if invoice.customer.customer_type == "Firma":
        replacements["Vorname"] = invoice.customer.company_name # Map company name to Vorname placeholder if that's how template works, or add Firmenname
        # Adjust based on template analysis if needed.
    
    if invoice.include_breakfast:
        replacements["AnzahlFrst"] = str(invoice.nights) # Assuming 1 breakfast per night per person? Or just nights?
        replacements["Frst"] = f"{invoice.breakfast_price:.2f}".replace(".", ",")
        replacements["Frst_ges"] = f"{invoice.breakfast_total:.2f}".replace(".", ",")
    else:
        replacements["AnzahlFrst"] = "0"
        replacements["Frst"] = "0,00"
        replacements["Frst_ges"] = "0,00"

    # Landlord data from UserProfile
    try:
        profile = invoice.user.profile
        landlord_address = f"{profile.company_name}\n{profile.street} {profile.house_number}\n{profile.postal_code} {profile.city}"
        if profile.phone:
            landlord_address += f"\nTel: {profile.phone}"
        if profile.email:
            landlord_address += f"\nEmail: {profile.email}"
        
        # Add bank details and tax number
        if profile.bank_details:
            replacements["Bankverbindung"] = profile.bank_details
        else:
            replacements["Bankverbindung"] = ""
            
        if profile.tax_number:
            replacements["Steuernummer"] = profile.tax_number
        else:
            replacements["Steuernummer"] = ""
    except Exception:
        # Fallback if no profile exists
        landlord_address = "Vermieter Adresse nicht konfiguriert"
        replacements["Bankverbindung"] = ""
        replacements["Steuernummer"] = ""

    replacements["{Vermieter_Anschrift}"] = landlord_address
    # Also support without braces if that's how it ended up (though script put braces)
    replacements["Vermieter_Anschrift"] = landlord_address

    replace_text_in_doc(doc, replacements)

    # Save temporary docx
    temp_docx = os.path.join(settings.MEDIA_ROOT, "invoices", f"temp_{invoice.invoice_number}.docx")
    os.makedirs(os.path.dirname(temp_docx), exist_ok=True)
    doc.save(temp_docx)

    # Save DOCX to model
    if os.path.exists(temp_docx):
        with open(temp_docx, "rb") as f:
            invoice.docx_file.save(f"Rechnung_{invoice.invoice_number}.docx", File(f), save=True)

    # Convert to PDF using xhtml2pdf
    try:
        from xhtml2pdf import pisa
        from django.template.loader import render_to_string
        
        # Determine greeting
        greeting = "Sehr geehrte Damen und Herren"
        c = invoice.customer
        if c.customer_type == 'Firma':
            greeting = "Sehr geehrte Damen und Herren"
        elif c.is_custom_salutation and c.custom_salutation:
            greeting = c.custom_salutation
        else:
            # Standard salutation logic
            title_part = f" {c.title}" if c.title else ""
            if c.salutation == "Herr":
                greeting = f"Sehr geehrter Herr{title_part} {c.last_name}"
            elif c.salutation == "Frau":
                greeting = f"Sehr geehrte Frau{title_part} {c.last_name}"
            else:
                # Fallback for Divers or empty
                greeting = f"Guten Tag{title_part} {c.first_name} {c.last_name}"

        # Prepare context for the HTML template
        context = {
            'invoice': invoice,
            'customer': invoice.customer,
            'landlord_info': '',
            'sender_line': '',
            'greeting': greeting,
            'company_name': '',
            'bank_details': '',
            'tax_number': '',
        }
        
        # Get landlord info from UserProfile
        try:
            profile = invoice.user.profile
            context['landlord_info'] = f"{profile.company_name}<br>{profile.street} {profile.house_number}<br>{profile.postal_code} {profile.city}"
            if profile.phone:
                context['landlord_info'] += f"<br>Tel: {profile.phone}"
            if profile.email:
                context['landlord_info'] += f"<br>Email: {profile.email}"
            
            context['sender_line'] = f"{profile.company_name}, {profile.street} {profile.house_number}, {profile.postal_code} {profile.city}"
            context['company_name'] = profile.company_name
            context['bank_details'] = profile.bank_details
            context['tax_number'] = profile.tax_number
        except Exception:
            pass
        
        # Render the HTML template
        html_string = render_to_string('invoices/invoice_pdf.html', context)
        
        # Generate PDF
        pdf_path = os.path.join(settings.MEDIA_ROOT, "invoices", f"Rechnung_{invoice.invoice_number}.pdf")
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        
        with open(pdf_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(html_string.encode('utf-8'), dest=pdf_file, encoding='utf-8')
        
        if not pisa_status.err and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                invoice.pdf_file.save(f"Rechnung_{invoice.invoice_number}.pdf", File(f), save=False)
            os.remove(pdf_path)
    except Exception as e:
        # Log error but don't fail completely if DOCX was saved
        print(f"PDF generation failed: {e}")
        
    invoice.save()
    
    # Cleanup temp docx
    if os.path.exists(temp_docx):
        os.remove(temp_docx)

